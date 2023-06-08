import docx
import io
import datetime
from .models import Client, Complects, Columns, BuilderObject, WaterConsumptionLevel
from docx.document import Document
from docx.table import Table
class DocxGenerator:
    def __init__(self) -> None:
        pass

    def _write_title_pge(self, client:Client):
        self.doc.add_heading("ПРОЕКТНАЯ СМЕТА")
        self.doc.add_heading("Комплексной системы очистки воды", level=2)
        self.doc.add_heading(f"№ <<НУЖНО ДОБАВИТЬ СИСТЕМУ НУМЕРАЦИИ СМЕТ>>", level=2)
        # self.doc.add_paragraph("Заказчик: <<ИНФОРМАЦИЯ О ЗАКАЗЧИКЕ>>")
        self.doc.add_paragraph(f"Ф.И.О. {client.first_name} {client.last_name}")
        self.doc.add_paragraph(f"Адрес ")
        self.doc.add_paragraph(f"Телефон {client.phone_number}")
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_heading("г Москва", level=3)
        curr_date = datetime.datetime.now().isoformat()
        curr_date = curr_date[0:curr_date.index('T')]
        self.doc.add_paragraph(curr_date)
        self.doc.add_page_break()
    
    def _write_begin_part(self):
        self.doc.add_heading("Уважаемый клиент!")
        self.doc.add_paragraph(
        '''Благодарим Вас за обращение в нашу компанию!
Предлагаем Вам на рассмотрение ТКП по поставке системы водоподготовки холодной воды по: 
механической очистке, осветлению, удалению соединений железа, умягчению, 
улучшению органолептических показателей
''')
    def _write_builder_info(self, builder_object:BuilderObject,water_consumption:WaterConsumptionLevel):
        self.doc.add_paragraph().add_run(text="1.	 Исходные данные")
        self.doc.add_paragraph().add_run(text="Краткое описание объекта")
        table:Table = self.doc.add_table(8,2)

        table.cell(0,0).text = 'Объект, где монтируется система водоподготовки'
        table.cell(0,1).text = builder_object.builder_type
        table.cell(1,0).add_paragraph('Место, выделяемое под монтаж системы водоподготовки')
        table.cell(1,1).add_paragraph(builder_object.montage_place)
        table.cell(2,0).add_paragraph('Основной источник водоснабжения объекта')
        table.cell(2,1).add_paragraph(builder_object.main_water_source)
        table.cell(3,0).add_paragraph('Тип канализации')
        table.cell(3,1).add_paragraph(builder_object.sewerage_type)
        table.cell(4,0).add_paragraph('Максимальный часовой расход воды, м3/ч')
        table.cell(4,1).add_paragraph(str(water_consumption.water_consumption))
        table.cell(5,0).add_paragraph('Количество проживающих, чел.')
        table.cell(5,1).add_paragraph(str(water_consumption.people_number))
        
        table.cell(6,0).add_paragraph('Норма водопотребления на 1 чел. в сутки, м3')
        table.cell(6,1).add_paragraph(str(water_consumption.human_daily_norm))
        table.cell(7,0).add_paragraph('Суточный расчетный расход воды м3/сут.	0,45')
        table.cell(7,1).add_paragraph(str(water_consumption.dayly_consumption))
    
    def generate_contract(
            self,
            complect_info:Complects, 
            full_water_analisys=None,client:Client=None,
            builder_object:BuilderObject=None,
            water_consumption:WaterConsumptionLevel=None):
        
        self.doc:Document = docx.Document()
        # doc.add_heading('ПРОЕКТНАЯ СМЕТА', 0)
        # doc.add_paragraph('Комплексной системы очистки воды')

        self._write_title_pge(client)
        self._write_begin_part()
        self._write_builder_info(builder_object,water_consumption)

        stream = io.BytesIO()
        
        self.doc.save(stream)
        print(stream)
        stream.seek(0)
        return stream